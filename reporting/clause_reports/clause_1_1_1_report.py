import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ─────────────────────────────────────────────
# COLOUR / STYLE CONSTANTS
# ─────────────────────────────────────────────
PURPLE       = RGBColor(75, 0, 130)   # #4B0082
WHITE        = RGBColor(255, 255, 255)
GREEN        = RGBColor(0, 128, 0)
RED          = RGBColor(255, 0, 0)
PURPLE_HEX   = "4B0082"
LIGHT_PURPLE = "F3ECFA"              # screenshot block background


class Clause111Report:

    def __init__(self, context, results):
        self.context = context
        self.results = results
        self.output_dir = context.evidence.run_dir

    # ─────────────────────────────────────────
    # HEADING  (purple text + purple underline)
    # ─────────────────────────────────────────
    def _add_itsar_heading(self, doc, text, level=1):
        """Styled section heading with a purple bottom border."""
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(14)
        para.paragraph_format.space_after  = Pt(2)

        run = para.add_run(text)
        run.bold            = True
        run.font.size       = Pt(16 if level == 1 else 14)
        run.font.color.rgb  = PURPLE

        # bottom border underline
        pPr   = para._p.get_or_add_pPr()
        pBdr  = OxmlElement("w:pBdr")
        bot   = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "12")
        bot.set(qn("w:space"), "2")
        bot.set(qn("w:color"), PURPLE_HEX)
        pBdr.append(bot)
        pPr.append(pBdr)

        return para

    # ─────────────────────────────────────────
    # SUB-HEADING  (purple text, no underline)
    # ─────────────────────────────────────────
    def _add_itsar_subheading(self, doc, text, level=2):
        """Lighter heading for numbered sub-sections."""
        para = doc.add_heading(text, level=level)
        run = para.runs[0]
        run.bold           = True
        run.font.size      = Pt(16 if level == 1 else 14)
        run.font.color.rgb = PURPLE
        para.paragraph_format.space_before = Pt(14)
        para.paragraph_format.space_after  = Pt(8)
        return para

    # ─────────────────────────────────────────
    # BOLD PARAGRAPH
    # ─────────────────────────────────────────
    def _add_bold_paragraph(self, doc, text):
        p = doc.add_paragraph()
        p.add_run(text).bold = True
        return p

    # ─────────────────────────────────────────
    # TABLE HEADER CELL  (purple bg, white bold)
    # ─────────────────────────────────────────
    def _style_table_header(self, cell, color=PURPLE_HEX):
        tcPr = cell._tc.get_or_add_tcPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:fill"), color)
        tcPr.append(shd)

        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold           = True
                run.font.color.rgb = WHITE

        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # cell padding
        cell.top_margin    = Inches(0.15)
        cell.bottom_margin = Inches(0.15)
        cell.left_margin   = Inches(0.15)
        cell.right_margin  = Inches(0.15)

    # ─────────────────────────────────────────
    # PREVENT TABLE ROW SPLIT
    # ─────────────────────────────────────────
    def _prevent_table_row_split(self, table):
        for row in table.rows:
            trPr      = row._tr.get_or_add_trPr()
            cantSplit = OxmlElement("w:cantSplit")
            trPr.append(cantSplit)

    # ─────────────────────────────────────────
    # KEEP-WITH-NEXT
    # ─────────────────────────────────────────
    def _keep_with_next(self, para):
        para.paragraph_format.keep_with_next = True
        para.paragraph_format.keep_together  = True
        return para

    # ─────────────────────────────────────────
    # PAGE NUMBER FOOTER
    # ─────────────────────────────────────────
    def _add_page_number(self, doc):
        footer    = doc.sections[0].footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run()
        for tag, text in [
            ("w:fldChar",   None),
            ("w:instrText", "PAGE"),
            ("w:fldChar",   None),
        ]:
            el = OxmlElement(tag)
            if tag == "w:fldChar":
                el.set(qn("w:fldCharType"), "begin" if text is None and not run._r.__len__() else "end")
            else:
                el.text = text
            run._r.append(el)

        # cleaner approach using three separate runs
        paragraph.clear()
        for fld_type in ("begin", None, "end"):
            r = paragraph.add_run()
            if fld_type in ("begin", "end"):
                fc = OxmlElement("w:fldChar")
                fc.set(qn("w:fldCharType"), fld_type)
                r._r.append(fc)
            else:
                it = OxmlElement("w:instrText")
                it.text = "PAGE"
                r._r.append(it)

    # ─────────────────────────────────────────
    # TITLE
    # ─────────────────────────────────────────
    def _add_title(self, doc, title_text="Telecom Compliance Automation Framework (TCAF)"):
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = title.add_run(title_text)
        run.bold           = True
        run.font.size      = Pt(26)
        run.font.color.rgb = PURPLE

        doc.add_paragraph()

    # ─────────────────────────────────────────
    # DUT DETAILS  (section 1)
    # ─────────────────────────────────────────
    def _add_dut_details(self, doc, context):
        self._add_itsar_heading(doc, "1. DUT DETAILS", 2)

        rows_data = [
            ("Device",           context.dut_model),
            ("Serial Number",    context.dut_serial),
            ("Firmware Version", context.dut_firmware),
            ("DUT IP Address",   context.ssh_ip),
        ]

        table = doc.add_table(rows=len(rows_data) + 1, cols=2)
        table.style = "Table Grid"

        # header row
        for i, header in enumerate(["Parameter", "Value"]):
            cell      = table.rows[0].cells[i]
            cell.text = header
            self._style_table_header(cell)

        # data rows
        for i, (key, val) in enumerate(rows_data, start=1):
            table.rows[i].cells[0].text = key
            table.rows[i].cells[1].text = str(val)

        self._add_data_cell_padding(table)
        self._prevent_table_row_split(table)

    # ─────────────────────────────────────────
    # ITSAR INFO  (section 2)
    # ─────────────────────────────────────────
    def _add_itsar_info(self, doc, context):
        self._add_itsar_heading(doc, "2. ITSAR INFORMATION", 2)

        table = doc.add_table(rows=3, cols=2)
        table.style = "Table Grid"

        for i, header in enumerate(["Field", "Value"]):
            cell      = table.rows[0].cells[i]
            cell.text = header
            self._style_table_header(cell)

        table.rows[1].cells[0].text = "ITSAR Section"
        table.rows[1].cells[1].text = context.itsar_section

        table.rows[2].cells[0].text = "Requirement"
        table.rows[2].cells[1].text = context.itsar_requirement

        self._add_data_cell_padding(table)
        self._prevent_table_row_split(table)

    # ─────────────────────────────────────────
    # REQUIREMENT DESCRIPTION  (section 3)
    # ─────────────────────────────────────────
    def _add_requirement(self, doc):
        self._add_itsar_heading(doc, "3. Requirement Description", 2)
        doc.add_paragraph(
            "The CPE shall communicate with authenticated management entities only. "
            "Protocols used for CPE management shall support mutual authentication "
            "mechanisms using authentication attributes such as username/password "
            "or equivalent mechanisms."
        )

    # ─────────────────────────────────────────
    # SCREENSHOT EVIDENCE BLOCK
    # ─────────────────────────────────────────
    def _add_screenshot_block(self, doc, title, image_path):
        """Lavender-background card with purple border and centred image."""
        TABLE_WIDTH = Inches(7.8)
        IMAGE_WIDTH = Inches(6.2)

        table = doc.add_table(rows=2, cols=1)
        table.alignment     = WD_TABLE_ALIGNMENT.CENTER
        table.allow_autofit = False

        self._prevent_table_row_split(table)

        # lock width
        for row in table.rows:
            cell       = row.cells[0]
            cell.width = TABLE_WIDTH

            # lavender background
            tcPr = cell._tc.get_or_add_tcPr()
            shd  = OxmlElement("w:shd")
            shd.set(qn("w:fill"), LIGHT_PURPLE)
            tcPr.append(shd)

            cell.top_margin    = Inches(0.2)
            cell.bottom_margin = Inches(0.2)
            cell.left_margin   = Inches(0.3)
            cell.right_margin  = Inches(0.3)

        # title cell
        title_cell = table.cell(0, 0)
        p_title    = title_cell.paragraphs[0]
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._keep_with_next(p_title)

        run = p_title.add_run(title)
        run.bold           = True
        run.font.size      = Pt(11)
        run.font.color.rgb = PURPLE

        # image cell
        img_cell = table.cell(1, 0)
        p_img    = img_cell.paragraphs[0]
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.keep_together = True
        p_img.add_run().add_picture(image_path, width=IMAGE_WIDTH)

        # purple border around the whole table
        tblPr      = table._tbl.tblPr
        tblBorders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right"):
            border = OxmlElement(f"w:{edge}")
            border.set(qn("w:val"),   "single")
            border.set(qn("w:sz"),    "12")
            border.set(qn("w:color"), PURPLE_HEX)
            tblBorders.append(border)
        tblPr.append(tblBorders)

        return table

    # ─────────────────────────────────────────
    # TEST EXECUTION  (section 4)
    # ─────────────────────────────────────────
    def _add_test_cases(self, doc, results):
        self._add_itsar_heading(doc, "4. Test Execution", 2)

        for idx, tc in enumerate(results, start=1):

            h = self._add_itsar_heading(doc, f"4.{idx} Test Case: {tc.name}", 2)
            self._keep_with_next(h)

            doc.add_paragraph(f"Description: {tc.description}")

            # coloured PASS / FAIL label
            p   = doc.add_paragraph("Result: ")
            run = p.add_run(tc.status)
            run.bold           = True
            run.font.color.rgb = GREEN if tc.status.upper() == "PASS" else RED

            # evidence screenshots
            for evidence in tc.evidence:
                screenshot = getattr(evidence, "screenshot", None)
                if screenshot and os.path.exists(screenshot):
                    self._add_screenshot_block(
                        doc,
                        f"Evidence Screenshot — {os.path.basename(screenshot)}",
                        evidence,
                    )

            # visual separator between test cases
            doc.add_paragraph()

    # ─────────────────────────────────────────
    # RESULT SUMMARY TABLE  (section 5)
    # ─────────────────────────────────────────
    def _add_result_summary(self, doc, results):
        h = self._add_itsar_heading(doc, "5. Test Case Result Summary", 2)
        self._keep_with_next(h)

        headers = ["SL No", "Test Case Name", "PASS/FAIL", "Remarks"]
        table   = doc.add_table(rows=len(results) + 1, cols=len(headers))
        table.style = "Table Grid"

        # header row
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            self._style_table_header(cell)

        # data rows
        for i, tc in enumerate(results, start=1):
            table.rows[i].cells[0].text = str(i)
            table.rows[i].cells[1].text = tc.name
            table.rows[i].cells[2].text = tc.status
            table.rows[i].cells[3].text = getattr(tc, "remarks", "")

            # colour the PASS/FAIL cell
            status_cell = table.rows[i].cells[2]
            for para in status_cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = (
                        GREEN if tc.status.upper() == "PASS" else RED
                    )

        self._add_data_cell_padding(table, skip_first_row=True)
        self._prevent_table_row_split(table)

    # ─────────────────────────────────────────
    # INTERNAL HELPER – cell padding for data rows
    # ─────────────────────────────────────────
    def _add_data_cell_padding(self, table, skip_first_row=True):
        start = 1 if skip_first_row else 0
        for row in table.rows[start:]:
            for cell in row.cells:
                cell.top_margin    = Inches(0.12)
                cell.bottom_margin = Inches(0.12)
                cell.left_margin   = Inches(0.12)
                cell.right_margin  = Inches(0.12)

    # ─────────────────────────────────────────
    # GENERATE REPORT  (main entry point)
    # ─────────────────────────────────────────
    def generate(self, context, results):
        report_path = os.path.join(self.output_dir, "tcaf_report.docx")

        doc = Document()

        self._add_page_number(doc)
        self._add_title(doc)

        self._add_dut_details(doc, context)
        doc.add_paragraph()

        self._add_itsar_info(doc, context)
        doc.add_paragraph()

        self._add_requirement(doc)
        doc.add_paragraph()

        self._add_test_cases(doc, results)
        doc.add_paragraph()

        self._add_result_summary(doc, results)

        doc.save(report_path)
        return report_path