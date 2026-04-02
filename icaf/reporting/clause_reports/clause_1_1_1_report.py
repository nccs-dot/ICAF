"""
reporting/clause_reports/clause_1_1_1_report.py
─────────────────────────────────────────────────────────────────────────────
Report generator for ITSAR 1.1.1 - Management Protocols Entity Mutual Auth.

Everything that was previously hard-coded is now resolved from three sources:

  1. context.profile  (ProfileLoader)
       DUT IP, credentials, SNMP user/passwords, SSH key paths, gRPC port,
       web login URL - all the per-device values the TCs use at runtime.
       The report builds the exact same commands the TCs ran, with the real
       DUT IP substituted in.

  2. icaf/reporting/report_config.yaml
       Document-level metadata that never changes per run:
       reviewer names, organisation name, revision history entries,
       "prepared for" text.  Edit this file to customise every report.

  3. context  (RuntimeContext)
       DUT name/version/hashes, ssh_ip, start/end times, run_dir.

Credentials are redacted to "[REDACTED]" by default.
Set  ICAF_SHOW_CREDENTIALS=1  to embed them verbatim (internal use only).

Status-aware OCR text  (PASS / FAIL / NOT RUN)
───────────────────────────────────────────────
YAML spec carries three variants per TC:
  observation_pass/fail/not_run  |  conclusion_pass/fail/not_run
  remarks_pass/fail/not_run
The generator picks the correct set automatically based on runtime status.

AI enrichment  (optional)
─────────────────────────
When ANTHROPIC_API_KEY is present and real evidence exists the AI enriches
the pre-authored YAML observation with specific command/output details.
─────────────────────────────────────────────────────────────────────────────
"""

from __future__ import annotations

import logging
import os
import re
import textwrap
from pathlib import Path
from typing import Any

import requests
import yaml

from docx.shared import RGBColor, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from icaf.reporting.helpers import (
    PURPLE, LIGHT_PURPLE, DARK_GREY, MID_GREY,
    TABLE_HEADER_BG, TABLE_ALT_BG, PASS_GREEN, FAIL_RED, WHITE,
    NOT_RUN_COLOR, NOT_RUN_BG,
    HEX_PURPLE, HEX_PASS_GREEN, HEX_FAIL_RED,
    _style_cell, _para_in_cell, _set_table_width, _set_col_widths,
    section_heading, sub_heading, tc_heading,
    body_para, label_value_para, bullet_item, numbered_item,
    spacer, terminal_block, add_screenshot, status_result_table,
    two_col_info_table,
    build_doc_with_header_footer,
)
from icaf.reporting.spec_loader import load_clause_spec

logger = logging.getLogger(__name__)

# ── Anthropic API ─────────────────────────────────────────────────────────────
_ANTHROPIC_API_URL = "https://api.anthropic.com/v1/messages"
_ANTHROPIC_MODEL   = "claude-sonnet-4-20250514"

# ── Status constants ──────────────────────────────────────────────────────────
_PASS    = "PASS"
_FAIL    = "FAIL"
_NOT_RUN = "NOT RUN"

# Show credentials verbatim?  Default: redact.
_SHOW_CREDS = os.environ.get("ICAF_SHOW_CREDENTIALS", "0") == "1"

# Status accent colours {status: (bg_hex, border_hex, RGBColor)} ────────────
_ACCENT = {
    _PASS:    ("E8F5E9", "006400", PASS_GREEN),
    _FAIL:    ("FFF3F3", "CC0000", FAIL_RED),
    _NOT_RUN: ("F5F5F5", "AAAAAA", NOT_RUN_COLOR),
}


# ─────────────────────────────────────────────────────────────────────────────
# report_config.yaml loader
# ─────────────────────────────────────────────────────────────────────────────

_REPORT_CONFIG_PATH = Path(__file__).resolve().parent.parent / "report_config.yaml"


def _load_report_config() -> dict:
    """
    Load icaf/reporting/report_config.yaml.
    Returns an empty dict if the file is missing (all callers have fallbacks).
    """
    if _REPORT_CONFIG_PATH.is_file():
        with open(_REPORT_CONFIG_PATH, encoding="utf-8") as fh:
            return yaml.safe_load(fh) or {}
    logger.warning(
        "report_config.yaml not found at %s - using built-in defaults",
        _REPORT_CONFIG_PATH,
    )
    return {}


def _cfg(config: dict, dotted_key: str, default: Any = "") -> Any:
    """Dot-path accessor for the report config dict."""
    node = config
    for part in dotted_key.split("."):
        if not isinstance(node, dict):
            return default
        node = node.get(part, {})
    return node if node != {} else default

def _build_oam_rows(context):
    """
    Build OAM table rows using:
    - Excel (raw_protocols)
    - Scan (verified_protocols)
    """

    oam = getattr(context, "oam_context", None)

    if not oam:
        return None  # fallback to YAML

    raw = set(p.lower() for p in oam.get("raw_protocols", []))
    verified = set(p.lower() for p in oam.get("verified_protocols", []))

    all_protocols = sorted(raw.union(verified))

    rows = []

    for proto in all_protocols:
        configured = proto in raw
        detected = proto in verified

        if configured and detected:
            status = "Yes (Verified)"
        elif configured and not detected:
            status = "Yes (Not Detected)"
        elif not configured and detected:
            status = "No (Unexpected)"
        else:
            status = "No"

        rows.append((proto.upper(), status))

    return rows


# ─────────────────────────────────────────────────────────────────────────────
# Dynamic command renderer
# Builds the exact commands the TCs ran using live profile + context values.
# ─────────────────────────────────────────────────────────────────────────────

def _redact(value: str) -> str:
    return value if _SHOW_CREDS else "[REDACTED]"


class CommandRenderer:
    """
    Builds per-TC CLI step lists with real DUT IP, usernames, and (optionally)
    passwords from context.profile and the RuntimeContext.

    Each public method returns list[str] - one string per step shown in the
    report's "b. Execution Steps" section.
    """

    def __init__(self, context: Any) -> None:
        self._ctx  = context
        prof = getattr(context, "profile", None)
        self._prof = prof  # may be ProfileLoader, dict, or None

    # ── profile accessor helpers ──────────────────────────────────────────

    def _p(self, key: str, default: Any = "") -> Any:
        """Read dotted key from profile; return default if profile absent."""
        if self._prof is None:
            return default
        if hasattr(self._prof, "get"):
            v = self._prof.get(key, default)
            return v if v is not None else default
        return default

    def _plist(self, key: str) -> list:
        if self._prof is None:
            return []
        if hasattr(self._prof, "get_list"):
            return self._prof.get_list(key) or []
        v = self._p(key, [])
        if isinstance(v, list):
            return v
        return [v] if v else []

    # ── context shortcuts ─────────────────────────────────────────────────

    @property
    def ip(self) -> str:
        return getattr(self._ctx, "ssh_ip", "") or ""

    @property
    def ssh_user(self) -> str:
        return getattr(self._ctx, "ssh_user", "") or ""

    @property
    def ssh_password(self) -> str:
        return getattr(self._ctx, "ssh_password", "") or ""

    def _ssh_target(self) -> str:
        tpl = self._p("ssh.target", "{user}@{ip}")
        return tpl.format(user=self.ssh_user, ip=self.ip)

    def _ssh_cmd(self) -> str:
        binary  = self._p("ssh.binary", "ssh")
        options = self._plist("ssh.connect_options")
        opts    = " ".join(str(o) for o in options if o)
        parts   = [binary] + ([opts] if opts else []) + [self._ssh_target()]
        return " ".join(parts)

    # ── TC1: SNMPv3 positive ──────────────────────────────────────────────

    def tc1_steps(self) -> list[str]:
        target    = self._p("snmp.target") or self.ip
        community = self._p("snmp.community", "public")
        user      = self._p("snmp.user",
                            getattr(self._ctx, "snmp_user", None) or "User1")
        auth_pass = self._p("snmp.auth_pass",
                            getattr(self._ctx, "snmp_auth_pass", None) or "")
        priv_pass = self._p("snmp.priv_pass",
                            getattr(self._ctx, "snmp_priv_pass", None) or "")
        oid       = "1.3.6.1.2.1.1.3.0"

        v1_cmd  = f"snmpget -v1  -c {community} {target} {oid}"
        v2_cmd  = f"snmpget -v2c -c {community} {target} {oid}"
        v3_cmd  = (
            f"snmpget -v3 -u {user} -l authPriv "
            f"-a SHA -A {_redact(auth_pass)} -x AES -X {_redact(priv_pass)} "
            f"{target} {oid}"
        )
        weak_cmd = (
            f"snmpget -v3 -u user2 -l authPriv "
            f"-a MD5 -A {_redact('weak_pass')} -x DES -X {_redact('weak_pass')} "
            f"{target} {oid}"
        )
        return [
            f"Verify SNMPv1 is disabled (expect timeout):  {v1_cmd}",
            "Capture SNMPv1 packets in Wireshark; confirm DUT sends no response.",
            f"Verify SNMPv2c is disabled (expect timeout):  {v2_cmd}",
            "Capture SNMPv2c packets in Wireshark; confirm DUT sends no response.",
            f"Log in to DUT via SSH:  {self._ssh_cmd()}",
            f"Authenticate: username={self.ssh_user}  password={_redact(self.ssh_password)}",
            f"Configure SNMPv3 group:  "
            f"snmp-agent group v3 SNMP_Group privacy",
            f"Configure SNMPv3 user ({user}):  "
            f"snmp-agent usm-user v3 {user} SNMP_GROUP simple "
            f"authentication-mode sha {_redact(auth_pass)} "
            f"privacy-mode aes128 {_redact(priv_pass)}",
            f"Run SNMPv3 GET with valid credentials (expect sysUpTime response):  {v3_cmd}",
            "Capture SNMPv3 traffic in Wireshark; confirm traffic is encrypted.",
            f"Attempt SNMPv3 GET with weak algorithms MD5/DES (expect rejection):  {weak_cmd}",
        ]

    # ── TC2: SNMPv3 invalid credentials ──────────────────────────────────

    def tc2_steps(self) -> list[str]:
        target        = self._p("snmp.target") or self.ip
        user          = self._p("snmp.user",
                                getattr(self._ctx, "snmp_user", None) or "User1")
        priv_pass     = self._p("snmp.priv_pass",
                                getattr(self._ctx, "snmp_priv_pass", None) or "")
        bad_auth_pass = self._p("snmp.bad_auth_pass", "WrongPass")
        oid           = "1.3.6.1.2.1.1.3.0"

        bad_cmd = (
            f"snmpget -v3 -u {user} -l authPriv "
            f"-a SHA -A {_redact(bad_auth_pass)} -x AES -X {_redact(priv_pass)} "
            f"{target} {oid}"
        )
        return [
            f"Attempt SNMPv3 GET with incorrect authentication password:  {bad_cmd}",
            "Observe that snmpget returns:  "
            "'Authentication failure (incorrect password, community or key)'.",
            "Capture packets in Wireshark; confirm authentication failure packet "
            "is encrypted and no MIB data is returned.",
        ]

    # ── TC3: SSH mutual auth (password) ──────────────────────────────────

    def tc3_steps(self) -> list[str]:
        bad_password = self._p("ssh.bad_password", "wrongpassword")
        nmap_cmd     = f"sudo nmap -p 22 --script ssh2-enum-algos {self.ip}"
        return [
            f"Verify SSH is reachable on the DUT:  {self._ssh_cmd()}",
            "Verify SSH server version via DUT command:  display ssh server status",
            f"Enumerate SSH cipher suites:  {nmap_cmd}",
            f"Positive case - authenticate with correct credentials: "
            f"username={self.ssh_user}  password={_redact(self.ssh_password)}; "
            "confirm session establishment.",
            "Capture SSH handshake in Wireshark; verify key-exchange "
            "(diffie-hellman-group14-sha256) and encryption (aes128-ctr) algorithms.",
            f"Negative case - attempt SSH login with wrong password "
            f"({_redact(bad_password)}); confirm 'Permission denied' is returned.",
            "Capture negative-case packets in Wireshark; verify traffic "
            "remains encrypted and no sensitive data is disclosed.",
        ]

    # ── TC4: SSH correct public key ───────────────────────────────────────

    def tc4_steps(self) -> list[str]:
        key_path = self._p("ssh.pubkey.key_path", "~/.ssh/id_ecdsaa")
        key_name = self._p("ssh.pubkey.dut_key_name", "PUBBKEY")
        dut_user = self._p("ssh.pubkey.dut_user", "Test5")
        key_file = key_path.split("/")[-1]
        login_cmd = (
            f"ssh -o IdentitiesOnly=yes -i {key_path} {dut_user}@{self.ip}"
        )
        return [
            f"Generate ECDSA-256 key pair on tester:  "
            f"ssh-keygen -t ecdsa -b 256 -f {key_path}",
            f"Confirm public key file:  cat {key_path}.pub",
            f"Transfer public key to DUT via SFTP:  "
            f"sftp {self.ssh_user}@{self.ip}  →  put {key_path}.pub /{key_file}.pub",
            f"Log in to DUT ({self._ssh_cmd()}) and import the public key:  "
            f"public-key peer {key_name} import sshkey flash:/{key_file}.pub",
            f"Create DUT local user {dut_user} with SSH service type and level-10 role.",
            f"Assign public key to {dut_user}:  "
            f"ssh user {dut_user} service-type all "
            f"authentication-type publickey assign publickey {key_name}",
            f"Login using the correct key (no password expected):  {login_cmd}",
            "Confirm successful login via:  display logbuffer  "
            "(expect 'Accepted publickey for {dut_user}')".format(dut_user=dut_user),
            "Capture SSH session in Wireshark; confirm encrypted key exchange.",
        ]

    # ── TC5: SSH wrong public key ─────────────────────────────────────────

    def tc5_steps(self) -> list[str]:
        wrong_key = self._p("ssh.pubkey.wrong_key_path", "~/.ssh/wrong_keyy")
        dut_user  = self._p("ssh.pubkey.dut_user", "Test5")
        login_cmd = (
            f"ssh -o IdentitiesOnly=yes -i {wrong_key} {dut_user}@{self.ip}"
        )
        return [
            f"Generate an unregistered ECDSA key pair:  "
            f"ssh-keygen -t ecdsa -b 256 -f {wrong_key}",
            f"Attempt SSH login using the unregistered key:  {login_cmd}",
            "Observe that DUT closes the connection with:  "
            "'Permission denied (publickey)'.",
            "Verify the failure in DUT log buffer:  display logbuffer",
            "Capture packets in Wireshark; confirm all traffic is encrypted "
            "and no management access is granted.",
        ]

    # ── TC6: HTTPS valid login ────────────────────────────────────────────

    def tc6_steps(self) -> list[str]:
        raw_url  = self._p("web.login_url", "https://{ip}/")
        login_url = raw_url.format(ip=self.ip) if "{ip}" in raw_url else raw_url
        username  = (
            self._p("web.username")
            or getattr(self._ctx, "web_username", None)
            or "admin"
        )
        password  = (
            self._p("web.password")
            or getattr(self._ctx, "web_password", None)
            or ""
        )
        tls10_raw = self._p("tls.tls10_test_command",
                            "openssl s_client -connect {ip}:443 -tls1")
        tls11_raw = self._p("tls.tls11_test_command",
                            "openssl s_client -connect {ip}:443 -tls1_1")
        tls10_cmd = tls10_raw.format(ip=self.ip)
        tls11_cmd = tls11_raw.format(ip=self.ip)
        return [
            f"Open web browser and navigate to the DUT login page:  {login_url}",
            f"Enter valid admin credentials - username: {username}  "
            f"password: {_redact(password)}",
            "Verify the management dashboard is displayed "
            "(System Information page confirming authenticated access).",
            "Capture HTTPS traffic in Wireshark; verify TLS Client Hello "
            "and Server Hello exchange.",
            "Confirm TLS version (1.2 or 1.3) and approved cipher suite "
            "from Wireshark capture.",
            f"Verify that TLS 1.0 is rejected by the DUT:  {tls10_cmd}",
            f"Verify that TLS 1.1 is rejected by the DUT:  {tls11_cmd}",
        ]

    # ── TC7: HTTPS invalid login ──────────────────────────────────────────

    def tc7_steps(self) -> list[str]:
        raw_url   = self._p("web.login_url", "https://{ip}/")
        login_url = raw_url.format(ip=self.ip) if "{ip}" in raw_url else raw_url
        username  = self._p("web.username", "admin")
        bad_pass  = self._p("web.bad_password", "WrongAdmin")
        return [
            f"Open web browser and navigate to:  {login_url}",
            f"Enter incorrect credentials - username: {username}  "
            f"password: {_redact(bad_pass)}",
            "Observe DUT displays 'Failed to log in' error and no "
            "management access is granted.",
            "Capture HTTPS traffic in Wireshark; confirm TLS encryption "
            "is maintained throughout the failed authentication exchange.",
        ]

    # ── TC8: gRPC/gNMI mutual auth ────────────────────────────────────────

    def tc8_steps(self) -> list[str]:
        port       = self._p("grpc.port", "50051")
        pki_domain = self._p("grpc.pki_domain", "grpc_pki")
        dut_user   = self._p("grpc.dut_user", "Test1")
        dut_pass   = self._p("grpc.dut_password", "")
        bad_pass   = self._p("grpc.bad_password", "WrongPass")
        cn         = self._p("grpc.dut_cn") or self.ip

        grpcurl_ok = (
            f"grpcurl -insecure -cert grpc.crt -key grpc.key -cacert ca.pem "
            f"-d '{{\"username\":\"{dut_user}\",\"password\":\"{_redact(dut_pass)}\"}}' "
            f"{self.ip}:{port} gnmi.gNMI/Login"
        )
        grpcurl_bad = (
            f"grpcurl -insecure -cert grpc.crt -key grpc.key -cacert ca.pem "
            f"-d '{{\"username\":\"{dut_user}\",\"password\":\"{_redact(bad_pass)}\"}}' "
            f"{self.ip}:{port} gnmi.gNMI/Login"
        )
        return [
            f"Generate CA key/cert, gRPC client key/cert, and PKCS#12 bundle "
            f"on tester (CN={cn}) using OpenSSL.",
            f"Transfer grpc.p12 and ca.pem to DUT:  "
            f"sftp {self.ssh_user}@{self.ip}",
            f"Import CA certificate into DUT PKI domain:  "
            f"pki import domain {pki_domain} pem ca filename ca.pem",
            f"Import client PKCS#12 into DUT PKI domain:  "
            f"pki import domain {pki_domain} p12 local filename grpc.p12",
            f"Verify certificate installation:  "
            f"display pki certificate domain {pki_domain} local",
            f"Enable gRPC on DUT:  grpc pki domain {pki_domain}  "
            f"→  grpc enable  →  grpc port {port}",
            f"Create DUT user {dut_user} with password {_redact(dut_pass)} "
            "and network-admin role.",
            f"Positive - run grpcurl with valid credentials; "
            f"verify token_id is returned:  {grpcurl_ok}",
            "Capture gRPC traffic in Wireshark; verify TLS handshake and "
            "encrypted application data.",
            f"Negative - run grpcurl with incorrect password; "
            f"verify no token_id is returned:  {grpcurl_bad}",
            "Capture negative-case traffic in Wireshark; confirm TLS "
            "encryption is maintained throughout.",
        ]

    # ── dispatcher ────────────────────────────────────────────────────────

    def steps_for(self, canonical: str) -> list[str]:
        """
        Return the dynamically-rendered step list for a TC canonical name.
        Returns [] if the TC is not mapped - caller falls back to YAML.
        """
        mapping = {
            "TC1_SNMPV3_POSITIVE":            self.tc1_steps,
            "TC2_SNMPV3_INVALID_CREDENTIALS": self.tc2_steps,
            "TC3_SSH_MUTUAL_AUTH":            self.tc3_steps,
            "TC4_SSH_CORRECT_PUBLIC_KEY":     self.tc4_steps,
            "TC5_SSH_INCORRECT_PUBLIC_KEY":   self.tc5_steps,
            "TC6_HTTPS_VALID_LOGIN":          self.tc6_steps,
            "TC7_HTTPS_INVALID_LOGIN":        self.tc7_steps,
            "TC8_GRPC_GNMI_MUTUAL_AUTH":      self.tc8_steps,
        }
        fn = mapping.get(canonical)
        if fn is None:
            return []
        try:
            return fn()
        except Exception as exc:  # noqa: BLE001
            logger.warning(
                "CommandRenderer.steps_for(%s) raised %s - falling back to YAML steps",
                canonical, exc,
            )
            return []


# ─────────────────────────────────────────────────────────────────────────────
# Generic helpers
# ─────────────────────────────────────────────────────────────────────────────

def _get(obj: Any, key: str, default: Any = None) -> Any:
    if isinstance(obj, dict):
        return obj.get(key, default)
    return getattr(obj, key, default)


def _clean_terminal_output(raw: str) -> list[str]:
    if not raw:
        return []
    ansi  = re.compile(r"\x1B(?:[@-Z\\-_]|\[[0-?]*[ -/]*[@-~])")
    text  = ansi.sub("", raw)
    lines = text.splitlines()
    while lines and not lines[0].strip():
        lines.pop(0)
    while lines and not lines[-1].strip():
        lines.pop()
    if len(lines) > 40:
        lines = lines[:40] + [f"... [{len(lines) - 40} more lines truncated]"]
    return lines


def _summarise_evidence_for_ai(evidence: list[dict]) -> str:
    parts: list[str] = []
    for i, ev in enumerate(evidence, start=1):
        cmd  = ev.get("command")
        out  = ev.get("output")
        shot = ev.get("screenshot")
        block: list[str] = [f"--- Evidence block {i} ---"]
        if cmd:
            block.append(f"Command: {cmd}")
        if out:
            lines = _clean_terminal_output(out)[:30]
            block.append("Output:\n" + "\n".join(lines))
        if shot:
            block.append(f"Screenshot: {str(shot).split('/')[-1]}")
        parts.append("\n".join(block))
    return "\n\n".join(parts) if parts else "(no evidence captured)"


# ─────────────────────────────────────────────────────────────────────────────
# Status-aware text selector
# ─────────────────────────────────────────────────────────────────────────────

def _pick_status_text(spec: dict, field: str, status: str, fallback: str = "") -> str:
    suffix_map = {_PASS: "pass", _FAIL: "fail", _NOT_RUN: "not_run"}
    suffix = suffix_map.get(status.upper(), "not_run")
    text = spec.get(f"{field}_{suffix}", "").strip()
    if text:
        return text
    text = spec.get(field, "").strip()   # legacy flat key
    if text:
        return text
    return fallback


def _default_observation(tc_name: str, status: str) -> str:
    if status == _PASS:
        return (
            f"The test case {tc_name} was executed and the DUT responded as expected. "
            "All verification steps completed successfully and the observed behaviour "
            "meets ITSAR 1.1.1 security requirements."
        )
    if status == _FAIL:
        return (
            f"The test case {tc_name} was executed but the DUT did not behave as required. "
            "One or more verification steps failed. The observed behaviour does not "
            "meet ITSAR 1.1.1 security requirements."
        )
    return "N/A"


def _default_conclusion(tc_name: str, status: str) -> str:
    if status == _PASS:
        return f"The DUT satisfied the security requirement for {tc_name} - COMPLIANT."
    if status == _FAIL:
        return f"The DUT did not satisfy the security requirement for {tc_name} - NON-COMPLIANT."
    return "N/A"


def _default_remark(tc_name: str, status: str) -> str:
    if status == _PASS:
        return f"The DUT correctly enforced the security requirement for {tc_name}."
    if status == _FAIL:
        return f"The DUT failed to satisfy the security requirement for {tc_name}."
    return "N/A - Test case was not executed."


# ─────────────────────────────────────────────────────────────────────────────
# AI observation enrichment (optional, ran TCs only)
# ─────────────────────────────────────────────────────────────────────────────

def _ai_enrich_observation(
    tc_name: str,
    spec: dict,
    result: dict,
    base_observation: str,
) -> str:
    api_key  = os.environ.get("ANTHROPIC_API_KEY", "")
    evidence = result.get("evidence") or []
    status   = result.get("status", "UNKNOWN").upper()

    has_real_evidence = any(
        ev.get("command") or ev.get("output") or ev.get("screenshot")
        for ev in evidence
    )
    if not has_real_evidence or not api_key:
        return base_observation

    evidence_text   = _summarise_evidence_for_ai(evidence)
    description     = spec.get("description", "").strip()
    expected_result = spec.get("expected_result", "").strip()

    system_prompt = textwrap.dedent("""
        You are a security test report writer for ITSAR evaluations.
        Enrich the base observation with specific details from the evidence
        (commands run, outputs received, Wireshark findings) while keeping
        the formal tone and 3–5 sentence length.
        Rules:
        - Use the base observation as your template.
        - Integrate only evidence details that genuinely add value.
        - Do NOT invent details not present in the evidence.
        - Use formal third-person technical language; no "I" or "we".
        - Respond with the enriched observation text ONLY - no JSON, no fences.
    """).strip()

    user_prompt = textwrap.dedent(f"""
        Test Case: {tc_name}
        Status: {status}

        Base Observation:
        {base_observation}

        Description:
        {description}

        Expected Result:
        {expected_result}

        Actual Evidence:
        {evidence_text}

        Write the enriched observation now.
    """).strip()

    try:
        resp = requests.post(
            _ANTHROPIC_API_URL,
            headers={
                "Content-Type":      "application/json",
                "x-api-key":         api_key,
                "anthropic-version": "2023-06-01",
            },
            json={
                "model":      _ANTHROPIC_MODEL,
                "max_tokens": 400,
                "system":     system_prompt,
                "messages":   [{"role": "user", "content": user_prompt}],
            },
            timeout=30,
        )
        resp.raise_for_status()
        enriched = resp.json()["content"][0]["text"].strip()
        if enriched:
            logger.debug("AI enrichment applied for %s", tc_name)
            return enriched
    except Exception as exc:  # noqa: BLE001
        logger.warning(
            "AI enrichment failed for %s (%s) - using base text", tc_name, exc
        )
    return base_observation


# ─────────────────────────────────────────────────────────────────────────────
# Main report class
# ─────────────────────────────────────────────────────────────────────────────

class Clause111Report:
    """Generates the full TCAF Word report for ITSAR clause 1.1.1."""

    def __init__(self, context: Any, results: list[dict]) -> None:
        self._ctx   = context
        self.results = list(results)
        self.spec   = load_clause_spec(self._ctx_get("clause", "1.1.1"))
        self._rcfg  = _load_report_config()    # report_config.yaml
        self._cmdr  = CommandRenderer(context) # dynamic command renderer

        self.output_dir = (
            self._ctx_get("run_dir")
            or getattr(self._ctx_get("evidence", None), "run_dir", None)
            or "output"
        )

        # ── Resolve DUT metadata from context ────────────────────────────
        raw_info = self._ctx_get("dut_info") or {}
        di = raw_info if isinstance(raw_info, dict) else vars(raw_info)

        dut_name = (
            di.get("dut_name")
            or self._ctx_get("dut_name")
            or self._ctx_get("dut_model", "Device Under Test")
        )
        dut_version = (
            di.get("dut_version")
            or self._ctx_get("dut_version")
            or self._ctx_get("dut_firmware", "N/A")
        )

        # Format timestamps (context.start_time may be a datetime object)
        raw_start = self._ctx_get("start_time")
        raw_end   = self._ctx_get("end_time")
        if raw_start is None:
            ctx_st    = getattr(context, "start_time", None)
            raw_start = (
                ctx_st.strftime("%d/%m/%Y")
                if hasattr(ctx_st, "strftime") else "N/A"
            )
        if raw_end is None:
            raw_end = raw_start

        # ── Resolve report_config.yaml metadata ──────────────────────────
        self.meta: dict[str, Any] = {
            # DUT / runtime
            "dut_name":     dut_name,
            "dut_version":  dut_version,
            "dut_ip":       getattr(context, "ssh_ip", "N/A") or "N/A",
            "os_hash":      di.get("os_hash")     or self._ctx_get("os_hash",     "N/A"),
            "config_hash":  di.get("config_hash") or self._ctx_get("config_hash", "N/A"),
            "start_time":   raw_start,
            "end_time":     raw_end,
            "itsar_id":     self._ctx_get("itsar_section", "1.1.1"),
            "itsar_version":"1.0.1",
            # from report_config.yaml
            "doc_number":   _cfg(self._rcfg, "organisation.doc_prefix",   "/IP ROUTER/1.1"),
            "created_by":   _cfg(self._rcfg, "document.created_by",       "Test Engineer"),
            "reviewed_by":  _cfg(self._rcfg, "document.reviewed_by",      "Reviewer"),
            "approved_by":  _cfg(self._rcfg, "document.approved_by",      "Approver"),
            "conducted_by": _cfg(self._rcfg, "document.conducted_by",     "Tester"),
            "reviewed_by2": _cfg(self._rcfg, "document.reviewed_by2",
                                 _cfg(self._rcfg, "document.reviewed_by", "Reviewer")),
            "prepared_for": _cfg(
                self._rcfg, "prepared_for.line",
                "National Centre of Communications Security, Bengaluru, "
                "Department of Telecom, Ministry of Communications, "
                "Government of India",
            ),
        }

        # ── TC lookup tables from YAML ────────────────────────────────────
        tc_specs = self.spec.get("testcases", {})
        self._canonical_ordered: list[str] = list(tc_specs.keys())

        self._runner_to_canonical: dict[str, str] = {}
        for canonical, tc_spec in tc_specs.items():
            runner_name = tc_spec.get("runner_name")
            if runner_name:
                self._runner_to_canonical[runner_name] = canonical
            self._runner_to_canonical[canonical] = canonical

        # ── Build result_map ──────────────────────────────────────────────
        self._result_map: dict[str, dict] = {}
        self._ran_canonical: list[str] = []

        for pos, r in enumerate(self.results):
            runner_name = _get(r, "name", f"UNKNOWN_TC{pos+1}")
            canonical   = self._resolve_canonical(runner_name, pos)
            ev_new      = _get(r, "evidence", []) or []

            if canonical in self._result_map:
                self._result_map[canonical]["evidence"].extend(ev_new)
                if _get(r, "status", "FAIL").upper() == "FAIL":
                    self._result_map[canonical]["status"] = "FAIL"
            else:
                self._result_map[canonical] = {
                    "name":        canonical,
                    "description": _get(r, "description", ""),
                    "status":      _get(r, "status", "FAIL").upper(),
                    "evidence":    list(ev_new),
                }
                self._ran_canonical.append(canonical)

        # ── Counters ─────────────────────────────────────────────────────
        ran_set = set(self._ran_canonical) & set(self._canonical_ordered)
        self._ran_count     = len(self._ran_canonical)
        self._total_defined = len(self._canonical_ordered)
        self._not_run_count = len(set(self._canonical_ordered) - ran_set)

        statuses = [self._result_map[n]["status"] for n in self._ran_canonical]
        self._pass_count = statuses.count("PASS")
        self._fail_count = len(statuses) - self._pass_count

        self.final_result = (
            "PASS"
            if self._fail_count == 0 and self._not_run_count == 0
            else "FAIL"
        )
        self.meta["final_result"] = self.final_result

        # ── Pre-compute OCR text ──────────────────────────────────────────
        tc_specs_dict = self.spec.get("testcases", {})
        self._tc_observation: dict[str, str] = {}
        self._tc_conclusion:  dict[str, str] = {}
        self._tc_remark:      dict[str, str] = {}

        for canonical in self._canonical_ordered:
            spec_tc = tc_specs_dict.get(canonical, {})

            if canonical in self._result_map:
                result = self._result_map[canonical]
                status = result["status"]
                obs = _pick_status_text(
                    spec_tc, "observation", status,
                    _default_observation(canonical, status),
                )
                obs = _ai_enrich_observation(canonical, spec_tc, result, obs)
                self._tc_observation[canonical] = obs
                self._tc_conclusion[canonical]  = _pick_status_text(
                    spec_tc, "conclusion", status,
                    _default_conclusion(canonical, status),
                )
                self._tc_remark[canonical] = _pick_status_text(
                    spec_tc, "remarks", status,
                    _default_remark(canonical, status),
                )
            else:
                self._tc_observation[canonical] = _pick_status_text(
                    spec_tc, "observation", _NOT_RUN, "N/A")
                self._tc_conclusion[canonical]  = _pick_status_text(
                    spec_tc, "conclusion", _NOT_RUN, "N/A")
                self._tc_remark[canonical] = _pick_status_text(
                    spec_tc, "remarks", _NOT_RUN,
                    "N/A - Test case was not executed in this evaluation run.")

    # ── helpers ───────────────────────────────────────────────────────────

    def _ctx_get(self, key: str, default: Any = None) -> Any:
        if isinstance(self._ctx, dict):
            return self._ctx.get(key, default)
        return getattr(self._ctx, key, default)

    def _resolve_canonical(self, runner_name: str, position: int) -> str:
        if runner_name in self._runner_to_canonical:
            return self._runner_to_canonical[runner_name]
        if position < len(self._canonical_ordered):
            logger.warning(
                "Runner name '%s' not in spec - using position fallback "
                "(slot %d → %s)",
                runner_name, position, self._canonical_ordered[position],
            )
            return self._canonical_ordered[position]
        logger.warning(
            "Runner name '%s' not in spec; position %d out of range",
            runner_name, position,
        )
        return runner_name

    # ── Section renderers ─────────────────────────────────────────────────

    def _section_revision_history(self, doc: Any) -> None:
        section_heading(doc, "Revision History")
        entries = _cfg(self._rcfg, "revision_history.entries", [])
        if not entries:
            entries = [
                {"version": "V.1.0", "date": "Initial Release",
                 "description": "NCCS Approved Test Plan with initial Test Cases."},
                {"version": "V.1.1", "date": "auto",
                 "description": "First Release of Test Report - automated evidence collected."},
            ]
        data_rows = []
        for e in entries:
            d = e.get("date", "")
            if d == "auto":
                d = self.meta["start_time"]
            data_rows.append((e.get("version", ""), d, e.get("description", "")))
        two_col_info_table(
            doc,
            headers    = ["Version", "Date",  "Changes"],
            col_widths = [1200,       1800,    6360],
            data_rows  = data_rows,
        )

    def _section_preface(self, doc: Any) -> None:
        section_heading(doc, (
            "TSTR for Evaluation of 1.1 Management Protocols Entity "
            "Mutual Authentication (1.1.1 of CSR)"
        ))
        sub_heading(doc, "Preface")
        label_value_para(doc, "DUT Details",   self.meta["dut_name"])
        label_value_para(doc, "DUT IP Address", self.meta["dut_ip"])
        spacer(doc)

        body_para(doc, "DUT Software Version:", bold=True)
        two_col_info_table(
            doc,
            headers    = ["Software Name",        "Software Version"],
            col_widths = [3600,                    5760],
            data_rows  = [("Device Firmware / OS", self.meta["dut_version"])],
        )
        spacer(doc)

        body_para(doc, "Digest Hash of OS:", bold=True)
        two_col_info_table(
            doc,
            headers    = ["Software Version",   "Hash Integrity Value"],
            col_widths = [3600,                   5760],
            data_rows  = [
                (self.meta["dut_version"], self.meta["os_hash"]),
                ("sshd_config",            self.meta["config_hash"]),
            ],
        )
        spacer(doc)

        body_para(doc, "Applicable ITSAR:", bold=True)
        for entry in self.spec["itsar"]["applicable_itsar"]:
            bullet_item(doc, f"{entry['ref']} ({entry['id']})")
        spacer(doc)
        body_para(doc, "ITSAR Version No.:", bold=True)
        for entry in self.spec["itsar"]["applicable_itsar"]:
            bullet_item(
                doc,
                f"{entry['version']} (Date of Release: {entry['release_date']})",
            )

    def _section_requirement(self, doc: Any) -> None:
        spacer(doc, large=True)
        itsar = self.spec["itsar"]
        body_para(doc,
            f"1. ITSAR Section No. & Name:  "
            f"{itsar['section_no']} {itsar['section_name']}",
            bold=True, color=PURPLE)
        body_para(doc,
            f"2. Security Requirement No. & Name:  "
            f"{itsar['requirement_no']} {itsar['requirement_name']}",
            bold=True, color=PURPLE)
        body_para(doc, "3. Requirement Description:", bold=True, color=PURPLE)
        body_para(doc, self.spec["requirement_description"].strip())

    def _section_dut_config(self, doc: Any) -> None:
        spacer(doc, large=True)
        body_para(doc, "4. DUT Configuration:", bold=True, color=PURPLE)
        body_para(doc, "Note: " + self.spec["dut_config"]["split_mode_note"].strip())
        spacer(doc)
        body_para(doc, "1) OAM Access supported by DUT:", bold=True)

        oam_rows = _build_oam_rows(self._ctx)

        if oam_rows:
            # Dynamic (Excel + Scan)
            two_col_info_table(
                doc,
                headers    = ["Protocol", "Supported"],
                col_widths = [2800,        6560],
                data_rows  = oam_rows,
            )
        else:
            # Fallback to YAML (existing behavior)
            two_col_info_table(
                doc,
                headers    = ["Protocol", "Supported"],
                col_widths = [2800,        6560],
                data_rows  = [
                    (row["protocol"], row["supported"])
                    for row in self.spec["dut_config"]["oam_access"]
                ],
            )
        body_para(doc, "Legend:", bold=True)
        bullet_item(doc, "Yes (Verified): Protocol configured and detected")
        bullet_item(doc, "Yes (Not Detected): Configured but not reachable")
        bullet_item(doc, "No (Unexpected): Running but not documented")
        spacer(doc)
        body_para(doc, "NOTE:", bold=True, color=FAIL_RED)
        body_para(doc, self.spec["dut_config"]["snmp_note"].strip())

    def _section_preconditions(self, doc: Any) -> None:
        spacer(doc, large=True)
        body_para(doc, "5. Pre-conditions:", bold=True, color=PURPLE)
        for cond in self.spec["preconditions"]:
            bullet_item(doc, cond)

    def _section_test_objective(self, doc: Any) -> None:
        spacer(doc, large=True)
        body_para(doc, "6. Test Objective:", bold=True, color=PURPLE)
        bullet_item(doc, self.spec["test_objective"].strip())

    def _section_test_plan(self, doc: Any) -> None:
        spacer(doc, large=True)
        body_para(doc, "7. Test Plan:", bold=True, color=PURPLE)
        bullet_item(doc, self.spec["test_plan"]["scope_note"].strip())
        spacer(doc)

        sub_heading(doc, "a. Number of Test Scenarios:")
        for idx, (_, tc) in enumerate(self.spec["testcases"].items(), start=1):
            numbered_item(doc, f"Test case {idx}: {tc['scenario']}")
        spacer(doc)

        sub_heading(doc, "b. Tools Required:")
        for tool in self.spec["test_plan"]["tools"]:
            bullet_item(doc, tool)
        spacer(doc)

        body_para(doc, "8. Expected Results for Pass:", bold=True, color=PURPLE)
        for idx, (_, tc) in enumerate(self.spec["testcases"].items(), start=1):
            numbered_item(doc, f"Test case {idx}: {tc['expected_result']}")
        spacer(doc)

        body_para(doc, "9. Expected Format of Evidence:", bold=True, color=PURPLE)
        body_para(doc, self.spec["test_plan"]["evidence_format"].strip())

    # ── Test execution ────────────────────────────────────────────────────

    def _section_test_execution(self, doc: Any) -> None:
        doc.add_page_break()
        section_heading(doc, "10. Test Execution")

        if self._not_run_count > 0:
            spacer(doc, small=True)
            body_para(
                doc,
                f"NOTE: This report covers {self._ran_count} of "
                f"{self._total_defined} defined test cases. "
                f"{self._not_run_count} case(s) were not executed in this run.",
                bold=True, color=FAIL_RED,
            )

        tc_specs     = self.spec.get("testcases", {})
        rendered_idx = 0

        for canonical in self._canonical_ordered:
            if canonical not in self._result_map:
                continue

            rendered_idx += 1
            result = self._result_map[canonical]
            spec   = tc_specs.get(canonical, {})
            status = result["status"]

            spacer(doc)
            tc_heading(doc, f"{rendered_idx}. Test Case Name: {canonical}")
            spacer(doc, small=True)

            # a. Description
            sub_heading(doc, "a. Test Case Description:")
            desc = (spec.get("description") or result.get("description") or "").strip()
            body_para(doc, desc or "No description available.")
            spacer(doc)

            # b. Execution steps - dynamic commands preferred, YAML fallback
            dynamic_steps = self._cmdr.steps_for(canonical)
            steps_to_render = dynamic_steps if dynamic_steps else spec.get("steps", [])
            if steps_to_render:
                sub_heading(doc, "b. Execution Steps:")
                for step in steps_to_render:
                    numbered_item(doc, step)
                spacer(doc)

            # c. Evidence
            evidence = result.get("evidence") or []
            has_ev   = any(
                ev.get("command") or ev.get("output") or ev.get("screenshot")
                for ev in evidence
            )
            if has_ev:
                sub_heading(doc, "c. Evidence Captured:")
                for ev in evidence:
                    print(ev)
                    command    = ev.get("command")
                    output_raw = ev.get("output")
                    screenshot = ev.get("screenshot")

                    if command:
                        label_value_para(doc, "Command Executed", command)
                    if output_raw:
                        lines = _clean_terminal_output(output_raw)
                        if lines:
                            body_para(doc, "Command Output:", bold=True)
                            terminal_block(doc, lines)
                            spacer(doc, small=True)
                    if screenshot:
                        from icaf.reporting.helpers import _resolve_screenshot_path
                        clean_path = _resolve_screenshot_path(screenshot)
                        if clean_path:
                            body_para(
                                doc,
                                f"Evidence Screenshot - "
                                f"{os.path.basename(clean_path)}",
                                bold=True,
                            )
                            add_screenshot(doc, clean_path, width_inches=5.5)
                            spacer(doc, small=True)
                        else:
                            fname = str(screenshot).split("/")[-1]
                            body_para(
                                doc,
                                f"Screenshot captured: {fname} "
                                "(file path recorded; attach manually if required)",
                                italic=True, color=MID_GREY,
                            )
                spacer(doc)

            # d. Test Observation
            sub_heading(doc, "d. Test Observation:")
            spacer(doc, small=True)
            body_para(doc, self._tc_observation.get(canonical, "N/A"))
            spacer(doc, small=True)

            # e. Conclusion
            sub_heading(doc, "e. Conclusion:")
            spacer(doc, small=True)
            body_para(doc, self._tc_conclusion.get(canonical, "N/A"))
            spacer(doc, small=True)

            # f. Remark
            sub_heading(doc, "f. Remark:")
            spacer(doc, small=True)
            body_para(doc, self._tc_remark.get(canonical, "N/A"))
            spacer(doc, small=True)

            # e. Evidence statement
            sub_heading(doc, "e. Evidence Provided:")
            body_para(
                doc,
                "Screenshots and command outputs are captured and attached "
                "during testing. Automated evidence is embedded above.",
            )
            spacer(doc, small=True)

            status_result_table(doc, result["status"])
            spacer(doc, large=True)
            _add_divider(doc)
            spacer(doc, large=True)

    # ── Result summary ────────────────────────────────────────────────────

    def _section_result_summary(self, doc: Any) -> None:
        doc.add_page_break()
        section_heading(doc, "11. Test Case Result Summary")
        spacer(doc)

        tc_specs   = self.spec.get("testcases", {})
        headers    = ["SL No.", "Test Case Name", "Result", "Remarks"]
        col_widths = [720, 3840, 1200, 3600]

        data_rows: list[tuple] = []
        for sl, canonical in enumerate(self._canonical_ordered, start=1):
            spec = tc_specs.get(canonical, {})
            if canonical in self._result_map:
                r      = self._result_map[canonical]
                status = r["status"]
                remarks = self._tc_remark.get(canonical, spec.get("remarks", ""))
            else:
                status  = _NOT_RUN
                remarks = self._tc_remark.get(
                    canonical, "Test case was not executed in this run.")
            data_rows.append((str(sl), canonical, status, remarks, status))

        totals_detail = (
            "All test cases passed successfully."
            if self.final_result == "PASS"
            else (
                f"{self._fail_count} case(s) failed"
                + (f", {self._not_run_count} not run." if self._not_run_count else ".")
            )
        )
        counts_str = f"{self._pass_count}P / {self._fail_count}F"
        if self._not_run_count:
            counts_str += f" / {self._not_run_count}NR"

        totals_row = (
            "",
            f"Total: {self._total_defined} defined  |  {self._ran_count} ran",
            counts_str,
            totals_detail,
            self.final_result,
        )
        _build_summary_table(doc, headers, col_widths, data_rows, totals_row)

    # ── Conclusion ────────────────────────────────────────────────────────

    def _section_conclusion(self, doc: Any) -> None:
        spacer(doc, large=True)
        section_heading(doc, "12. Test Conclusion")
        spacer(doc)

        for b in self.spec.get("conclusion_bullets", []):
            bullet_item(doc, b)

        if self._not_run_count > 0:
            bullet_item(
                doc,
                f"NOTE: This run executed {self._ran_count} of "
                f"{self._total_defined} defined test cases. The remaining "
                f"{self._not_run_count} case(s) were not run and must be "
                "completed before a final pass verdict can be issued.",
                bold=True,
            )

        spacer(doc)
        status_result_table(
            doc,
            status=self.final_result,
            label="Overall Evaluation Result",
            detail=(
                f"{self._pass_count} of {self._total_defined} cases passed."
                if not self._not_run_count
                else (
                    f"{self._pass_count} passed, "
                    f"{self._fail_count} failed, "
                    f"{self._not_run_count} not run."
                )
            ),
            wide=True,
        )

    # ── Entry point ───────────────────────────────────────────────────────

    def generate(self) -> str:
        os.makedirs(self.output_dir, exist_ok=True)
        report_path = os.path.join(self.output_dir, "tcaf_report.docx")

        doc = build_doc_with_header_footer(
            dut_name    = self.meta["dut_name"],
            dut_version = self.meta["dut_version"],
        )

        _build_front_page_dynamic(doc, self.meta)
        doc.add_page_break()

        self._section_revision_history(doc)
        doc.add_page_break()

        self._section_preface(doc)
        self._section_requirement(doc)
        self._section_dut_config(doc)
        self._section_preconditions(doc)
        self._section_test_objective(doc)
        self._section_test_plan(doc)

        self._section_test_execution(doc)
        self._section_result_summary(doc)
        self._section_conclusion(doc)

        doc.save(report_path)
        logger.info("Report saved: %s", report_path)
        return report_path


# ─────────────────────────────────────────────────────────────────────────────
# Dynamic front page - uses report_config.yaml metadata
# ─────────────────────────────────────────────────────────────────────────────

def _build_front_page_dynamic(doc: Any, meta: dict) -> None:
    """
    Front page built entirely from the enriched meta dict.
    Reviewer names, approver, doc number all come from report_config.yaml
    via the meta dict - nothing is hard-coded here.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from icaf.reporting.helpers import (
        spacer, four_col_header_table, _add_para_border_bottom,
        MID_GREY, PURPLE,
    )

    dut_name     = meta.get("dut_name",     "Device Under Test")
    start_time   = meta.get("start_time",   "N/A")
    end_time     = meta.get("end_time",     "N/A")
    doc_number   = meta.get("doc_number",   "N/A")
    created_by   = meta.get("created_by",   "Test Engineer")
    reviewed_by  = meta.get("reviewed_by",  "Reviewer")
    approved_by  = meta.get("approved_by",  "Approver")
    conducted_by = meta.get("conducted_by", "Tester")
    reviewed_by2 = meta.get("reviewed_by2", reviewed_by)
    prepared_for = meta.get("prepared_for", "")

    spacer(doc, large=True)
    spacer(doc, large=True)

    def _centered(text, size_pt, color, bold=False):
        p   = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold           = bold
        run.font.size      = Pt(size_pt)
        run.font.name      = "Arial"
        run.font.color.rgb = color
        return p

    _centered("TSTL Evaluation Test Report", 26, PURPLE, bold=True)
    _centered("for", 18, MID_GREY)
    p = _centered(dut_name, 24, PURPLE, bold=True)
    p.paragraph_format.space_after = Pt(14)
    _add_para_border_bottom(p, "4B0082", size=12)

    spacer(doc, large=True)

    if prepared_for:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("Prepared For:  ")
        r.bold             = True
        r.font.size        = Pt(12)
        r.font.name        = "Arial"
        r.font.color.rgb   = MID_GREY
        r2 = p.add_run(prepared_for)
        r2.font.size       = Pt(10)
        r2.font.name       = "Arial"
        r2.font.color.rgb  = MID_GREY

    spacer(doc, large=True)

    four_col_header_table(
        doc,
        headers    = ["Document No.", "Created By",  "Reviewed By", "Approved By"],
        data_rows  = [(doc_number,    created_by,     reviewed_by,   approved_by)],
        col_widths = [2340, 2340, 2340, 2340],
    )
    spacer(doc, small=True)
    four_col_header_table(
        doc,
        headers    = ["Test conducted by", "Test conducted on",
                      "Test reviewed by",  "Test reviewed on"],
        data_rows  = [(conducted_by, start_time, reviewed_by2, end_time)],
        col_widths = [2340, 2340, 2340, 2340],
    )


# ─────────────────────────────────────────────────────────────────────────────
# Module-private helpers
# ─────────────────────────────────────────────────────────────────────────────

def _add_divider(doc: Any) -> None:
    """Thin grey horizontal rule between test case blocks."""
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:color"), "DDDDDD")
    pBdr.append(bot)
    pPr.append(pBdr)


def _build_summary_table(
    doc:        Any,
    headers:    list[str],
    col_widths: list[int],
    data_rows:  list[tuple],
    totals_row: tuple,
) -> None:
    from docx.enum.table import WD_TABLE_ALIGNMENT

    table = doc.add_table(rows=0, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_width(table, sum(col_widths))
    _set_col_widths(table, col_widths)

    # Header
    hdr = table.add_row()
    for ci, (h, w) in enumerate(zip(headers, col_widths)):
        c = hdr.cells[ci]
        _style_cell(c, TABLE_HEADER_BG, HEX_PURPLE, w)
        _para_in_cell(c, h, bold=True, color=WHITE, center=True)

    # Data rows
    for ri, (sl, tc_name, status_text, remarks, status_key) in enumerate(data_rows):
        row    = table.add_row()
        row_bg = TABLE_ALT_BG if ri % 2 else "FFFFFF"
        sk     = status_key.upper()

        if sk == _PASS:
            s_bg  = "E8F5E9"
            s_col = RGBColor(0x00, 0x64, 0x00)
        elif sk == _NOT_RUN:
            s_bg  = NOT_RUN_BG
            s_col = NOT_RUN_COLOR
        else:
            s_bg  = "FFEBEE"
            s_col = RGBColor(0xCC, 0x00, 0x00)

        for ci, (val, w) in enumerate(
            zip([sl, tc_name, status_text, remarks], col_widths)
        ):
            c = row.cells[ci]
            if ci == 2:
                _style_cell(c, s_bg, "CCCCCC", w)
                _para_in_cell(c, val, bold=True, color=s_col, center=True)
            else:
                _style_cell(c, row_bg, "CCCCCC", w)
                _para_in_cell(c, val, color=DARK_GREY, center=(ci == 0))

    # Totals row
    tot     = table.add_row()
    fk      = totals_row[4].upper()
    tot_col = (
        RGBColor(0x00, 0x64, 0x00) if fk == _PASS
        else RGBColor(0xCC, 0x00, 0x00)
    )
    for ci, (val, w) in enumerate(zip(totals_row[:4], col_widths)):
        c = tot.cells[ci]
        _style_cell(c, LIGHT_PURPLE, "CCCCCC", w)
        _para_in_cell(c, val, bold=True, color=tot_col, center=(ci in (0, 2)))
